"""
LifeGate CEO LinkedIn Post Calendar – Week 1 Generator
Produces two Word documents:
  1. lifegate_linkedin_calendar_week1.docx  – calendar + prompts
  2. lifegate_linkedin_posts_week1.docx     – ready-to-publish posts
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ──────────────────────────────────────────────
# BRAND COLOURS
# ──────────────────────────────────────────────
TEAL   = RGBColor(0x0A, 0xAD, 0xA2)   # #0AADA2 – LifeGate primary
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF5, 0xF7, 0xFA)
MGREY  = RGBColor(0x64, 0x74, 0x8B)
GOLD   = RGBColor(0xF5, 0xA6, 0x23)   # accent for highlights

OUTPUT_DIR = "/workspaces/lifegate-mobile-app/public/docs"

# ──────────────────────────────────────────────
# WEEK 1 DATA
# ──────────────────────────────────────────────
WEEK1 = [
    {
        "day": "Monday, June 1, 2026",
        "post_number": "01",
        "theme": "Platform Mission & Vision",
        "domain": "Brand / Mission",
        "tone": "Visionary · Inspiring",
        "hashtags": "#DigitalHealth #HealthcareForAll #LifeGate #AIHealth #Nigeria",
        "cta": "Follow LifeGate for weekly updates.",
        "prompt": (
            "Write a captivating LinkedIn post from the perspective of the CEO of LifeGate, "
            "a Nigerian AI-assisted digital health platform. Open with a bold, emotionally resonant "
            "hook about the state of healthcare access in Nigeria — long queues, expensive visits, "
            "missed early diagnoses. Then introduce LifeGate as the solution: an always-available, "
            "AI-powered health companion backed by licensed physicians. Highlight the tagline "
            "'Your trusted digital health companion.' Close with a passionate call to action inviting "
            "followers to be part of the movement to democratise healthcare access. Tone: visionary, "
            "warm, confident. Length: 180–220 words. Include 5 relevant hashtags."
        ),
        "post": (
            "I grew up watching people in my community avoid hospitals — not because they didn't "
            "need care, but because the journey itself felt impossible.\n\n"
            "Long queues. Expensive consultations. Clinics hours away. And by the time someone finally "
            "got seen, a simple, treatable condition had become something far more serious.\n\n"
            "That pain point became the founding question behind LifeGate:\n\n"
            "👉 What if every person had a trusted health companion available 24/7 — one that listens, "
            "triages intelligently, and connects them to a real physician when it matters most?\n\n"
            "Today, LifeGate is that companion.\n\n"
            "Powered by cutting-edge AI and validated by licensed physicians, LifeGate bridges the gap "
            "between knowing something is wrong and actually getting help. We don't replace doctors — "
            "we make sure patients reach the right doctor, at the right time, every time.\n\n"
            "Our mission is simple: healthcare access for every Nigerian, regardless of zip code or "
            "income level.\n\n"
            "We're just getting started — and I'm proud of what this team has built.\n\n"
            "Follow along. This journey is going to change lives.\n\n"
            "#DigitalHealth #HealthcareForAll #LifeGate #AIHealth #Nigeria"
        ),
    },
    {
        "day": "Tuesday, June 2, 2026",
        "post_number": "02",
        "theme": "EDIS — AI-Powered Triage Engine",
        "domain": "Early Detection Intelligence System (EDIS)",
        "tone": "Educational · Exciting",
        "hashtags": "#AITriage #EarlyDetection #HealthTech #EDIS #LifeGate",
        "cta": "Comment 'EDIS' to learn more about our triage technology.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate explaining the EDIS (Early Detection "
            "Intelligence System) — the AI probabilistic triage engine at the heart of the platform. "
            "Use an analogy to make AI triage relatable (e.g., a highly experienced diagnostician in "
            "your pocket). Explain that EDIS analyses patient-reported symptoms conversationally, "
            "assigns risk probabilities, and knows exactly when to escalate to a licensed physician. "
            "Emphasise that EDIS also interprets vision and audiometry sensor test results. Highlight "
            "the dual-mode operation: General Health vs Clinical Diagnosis mode. Keep it accessible — "
            "no jargon. Tone: educational, enthusiastic, approachable. Length: 180–220 words."
        ),
        "post": (
            "What if you could carry a world-class diagnostician in your pocket?\n\n"
            "That's exactly what our EDIS — Early Detection Intelligence System — is designed to be.\n\n"
            "Here's how it works:\n\n"
            "🧠 A patient opens LifeGate and describes how they feel — in plain language.\n"
            "📊 EDIS analyses the symptoms conversationally, building a probabilistic risk profile in "
            "real time.\n"
            "🚨 When clinical risk is detected, it escalates immediately to a licensed physician — no "
            "delays, no guesswork.\n\n"
            "EDIS operates in two modes:\n"
            "→ General Health Mode — for everyday wellness questions\n"
            "→ Clinical Diagnosis Mode — for structured, physician-bound assessments\n\n"
            "And it goes beyond symptoms. EDIS can interpret vision screening and audiometry sensor "
            "test results, making it one of the most comprehensive AI triage engines built for the "
            "African healthcare context.\n\n"
            "The result? Faster identification of serious conditions. Fewer missed diagnoses. "
            "And a smarter path to the right care.\n\n"
            "Early detection saves lives. EDIS makes it accessible to everyone.\n\n"
            "💬 Comment 'EDIS' below to learn more about our technology.\n\n"
            "#AITriage #EarlyDetection #HealthTech #EDIS #LifeGate"
        ),
    },
    {
        "day": "Wednesday, June 3, 2026",
        "post_number": "03",
        "theme": "Physician Validation — Human Oversight in the Loop",
        "domain": "Physician Review & Validation",
        "tone": "Trust-building · Professional",
        "hashtags": "#PhysicianLed #MedicalTrust #AIandDoctors #LifeGate #SafeHealth",
        "cta": "Are you a licensed physician? DM us about joining our network.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate on why LifeGate is NOT just another "
            "AI health chatbot. Emphasise the mandatory physician validation layer — every clinical "
            "output generated by EDIS must be reviewed and validated by a licensed physician before "
            "reaching the patient. Describe the three-state physician workflow (Pending → Active → "
            "Completed) and the decision options (Approve, Edit with notes, Reject with reason). "
            "Mention MDCN licence verification as the gold standard for physician onboarding. "
            "Use the phrase 'AI does the groundwork; physicians own the outcome.' Build trust with "
            "medical professionals and patients alike. Tone: authoritative, reassuring. Length: "
            "180–220 words."
        ),
        "post": (
            "Let me be direct about something.\n\n"
            "LifeGate is NOT a replacement for doctors. It never will be.\n\n"
            "Here's what makes us fundamentally different from every other health AI tool in the "
            "market:\n\n"
            "🔒 EVERY clinical output on LifeGate is reviewed and validated by a licensed physician "
            "before it reaches a patient. Full stop.\n\n"
            "Our physician workflow is deliberate and rigorous:\n"
            "→ Pending: Case is triaged and queued\n"
            "→ Active: Physician reviews AI output, patient history & OLDCARTS HPI\n"
            "→ Completed: Physician approves, edits with notes, or rejects with documented reason\n\n"
            "Every physician on our platform is verified through MDCN — Nigeria's Medical and Dental "
            "Council — before they see a single case.\n\n"
            "Our philosophy: AI does the groundwork. Physicians own the outcome.\n\n"
            "This isn't a compromise. It's a clinical imperative. We built LifeGate this way because "
            "we believe that in healthcare, accountability must always rest with a human being.\n\n"
            "Patients deserve that standard. Our physician partners uphold it every day.\n\n"
            "Are you a licensed physician interested in joining our network? DM me directly.\n\n"
            "#PhysicianLed #MedicalTrust #AIandDoctors #LifeGate #SafeHealth"
        ),
    },
    {
        "day": "Thursday, June 4, 2026",
        "post_number": "04",
        "theme": "SafeMed Practices — Privacy, Consent & Data Safety",
        "domain": "SafeMed / NDPA Compliance / Audit",
        "tone": "Reassuring · Transparent",
        "hashtags": "#SafeMed #DataPrivacy #NDPA #PatientSafety #LifeGate",
        "cta": "Your health data is safe with LifeGate. Ask us how.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate introducing LifeGate's SafeMed "
            "practices — the platform's framework for patient data safety, privacy, and ethical "
            "AI in healthcare. Cover: full NDPA 2023 (Nigeria Data Protection Act) compliance, "
            "explicit consent flows before every data collection, complete audit logging of all "
            "system actions, anonymised epidemiological surveillance data for public health, "
            "and JWT-secured, role-based access control. Use a reassuring, transparent tone that "
            "builds patient trust. Mention that healthcare data is among the most sensitive "
            "personal data and LifeGate treats it as such. Avoid technical jargon. "
            "Length: 180–220 words."
        ),
        "post": (
            "When you share your health symptoms, your medical history, your fears about what "
            "might be wrong — you are trusting us with something deeply personal.\n\n"
            "We take that trust seriously. Seriously enough to build our entire platform around it.\n\n"
            "Introducing LifeGate SafeMed Practices — our framework for keeping your health data "
            "safe, private, and ethically managed:\n\n"
            "✅ Full NDPA 2023 Compliance — Every data operation aligns with Nigeria's Data "
            "Protection Act\n"
            "✅ Explicit Consent Flows — You control what data is collected, every step of the way\n"
            "✅ Complete Audit Logging — Every system action is logged for accountability and "
            "transparency\n"
            "✅ Anonymised Public Health Data — Epidemiological insights shared publicly without "
            "compromising individual identity\n"
            "✅ Role-Based Access Control — Only the right people see the right information\n\n"
            "Health data is not like other data. A breach isn't just an inconvenience — it can "
            "destroy trust, harm patients, and undermine care.\n\n"
            "We will never sell your data. We will never cut corners on compliance.\n\n"
            "SafeMed isn't a feature. It's a promise.\n\n"
            "#SafeMed #DataPrivacy #NDPA #PatientSafety #LifeGate"
        ),
    },
    {
        "day": "Friday, June 5, 2026",
        "post_number": "05",
        "theme": "Customer Success — Real Impact Stories",
        "domain": "Customer Success / Patient Experience",
        "tone": "Emotional · Celebratory",
        "hashtags": "#CustomerSuccess #HealthcareHero #LifeGate #PatientFirst #DigitalHealth",
        "cta": "Share this post with someone who deserves better healthcare access.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate sharing a compelling customer success "
            "story (fictional but realistic for the Nigerian context). The story should follow a "
            "patient — e.g., a 34-year-old teacher in Kano — who used LifeGate to identify early "
            "warning signs of hypertension through the AI chat, was escalated to a physician who "
            "confirmed the diagnosis, and received a personalised care plan through the app. "
            "Highlight the emotional arc: fear → accessible help → diagnosis → relief → action. "
            "Tie the story back to LifeGate's mission of democratising healthcare access. "
            "End with an emotional call to action asking followers to share the post. "
            "Tone: warm, personal, emotionally compelling. Length: 200–240 words."
        ),
        "post": (
            "Her name is Amina. She's 34 years old, a primary school teacher in Kano.\n\n"
            "For weeks she had been getting headaches. Dizziness in the afternoon. Sometimes her "
            "vision would blur during class. She told herself it was stress. She couldn't afford to "
            "take time off. The clinic was 45 minutes away, and the queue — well, you know the queue.\n\n"
            "One evening, she opened LifeGate.\n\n"
            "She typed her symptoms into the chat — headaches, dizziness, fatigue. The AI asked "
            "follow-up questions: family history, diet, how long had this been happening? Within "
            "minutes, EDIS flagged elevated hypertension risk and escalated her case to a physician.\n\n"
            "By the next morning, Dr. Okafor had reviewed her full profile, confirmed early-stage "
            "hypertension, and sent her a personalised care plan — including dietary changes, "
            "a referral to a local clinic for blood pressure monitoring, and a follow-up reminder "
            "set for two weeks later.\n\n"
            "Amina didn't have to take a day off. She didn't have to sit in a queue.\n\n"
            "She got answers. She got a plan. She got her health back.\n\n"
            "This is why we built LifeGate. For every Amina who deserves better access to care.\n\n"
            "If you know someone who needs this, please share this post. It could save a life. 💙\n\n"
            "#CustomerSuccess #HealthcareHero #LifeGate #PatientFirst #DigitalHealth"
        ),
    },
]


# ──────────────────────────────────────────────
# HELPER UTILITIES
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_run(para, text, bold=False, italic=False,
            size=11, color=None, font_name="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run


def heading(doc, text, level=1, color=DARK, size=None, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sizes = {1: 22, 2: 16, 3: 13, 4: 11}
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def body(doc, text, italic=False, color=DARK, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def divider(doc, color_hex="0AADA2"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────
# DOCUMENT 1 – CALENDAR & PROMPTS
# ──────────────────────────────────────────────

def build_calendar_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_para.paragraph_format.space_before = Pt(40)
    add_run(cover_para, "LifeGate", bold=True, size=36, color=TEAL)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, "CEO LinkedIn Post Calendar", bold=True, size=22, color=DARK)

    week_line = doc.add_paragraph()
    week_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(week_line, "WEEK 1  ·  June 1 – June 5, 2026", bold=False, size=14, color=MGREY)

    divider(doc)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tagline, '"Your trusted digital health companion."', italic=True, size=12, color=MGREY)

    doc.add_paragraph()

    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared for", "CEO, LifeGate (DHSub)"),
        ("Campaign Period", "Week 1 of 6 — June 2026"),
        ("Posts this week", "5 posts (Mon – Fri)"),
    ]
    for i, (label, value) in enumerate(meta_data):
        lc = meta_table.rows[i].cells[0]
        vc = meta_table.rows[i].cells[1]
        set_cell_bg(lc, "0AADA2")
        lc.paragraphs[0].clear()
        add_run(lc.paragraphs[0], label, bold=True, color=WHITE, size=10)
        vc.paragraphs[0].clear()
        add_run(vc.paragraphs[0], value, size=10, color=DARK)

    page_break(doc)

    # ── Introduction ──────────────────────────
    heading(doc, "About This Calendar", level=2, color=TEAL)
    divider(doc)
    body(doc,
         "This document contains the Week 1 LinkedIn Post Calendar for the CEO of LifeGate. "
         "Each entry includes the post theme, targeted functional domain, recommended tone, "
         "suggested hashtags, a call-to-action, and a detailed writing prompt. Use these prompts "
         "with any AI writing assistant (or your content team) to generate the final post. "
         "The companion document — LifeGate LinkedIn Posts Week 1 — contains the fully written, "
         "ready-to-publish versions of each post.")
    doc.add_paragraph()

    # ── Strategy overview table ────────────────
    heading(doc, "Week 1 Content Strategy Overview", level=2, color=TEAL)
    divider(doc)

    strat_table = doc.add_table(rows=6, cols=4)
    strat_table.style = "Table Grid"

    headers = ["Day", "Theme", "Domain", "Tone"]
    hrow = strat_table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        set_cell_bg(cell, "0AADA2")
        cell.paragraphs[0].clear()
        add_run(cell.paragraphs[0], h, bold=True, color=WHITE, size=10)

    for i, post in enumerate(WEEK1):
        row = strat_table.rows[i + 1]
        vals = [post["day"].split(",")[0], post["theme"], post["domain"], post["tone"]]
        for j, val in enumerate(vals):
            row.cells[j].paragraphs[0].clear()
            add_run(row.cells[j].paragraphs[0], val, size=9, color=DARK)
            if i % 2 == 0:
                set_cell_bg(row.cells[j], "F5F7FA")

    doc.add_paragraph()
    page_break(doc)

    # ── Individual post entries ────────────────
    heading(doc, "Week 1 — Post Prompts", level=1, color=TEAL, size=18, center=True)
    doc.add_paragraph()

    for idx, post in enumerate(WEEK1):
        # Post header bar (using a 1-row table for colour block)
        hdr_table = doc.add_table(rows=1, cols=1)
        hdr_table.style = "Table Grid"
        hdr_cell = hdr_table.rows[0].cells[0]
        set_cell_bg(hdr_cell, "0AADA2")
        hdr_cell.paragraphs[0].clear()
        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(hdr_cell.paragraphs[0],
                f"POST {post['post_number']}  ·  {post['day']}",
                bold=True, color=WHITE, size=12)

        # Detail table
        detail_table = doc.add_table(rows=5, cols=2)
        detail_table.style = "Table Grid"
        detail_data = [
            ("Theme",    post["theme"]),
            ("Domain",   post["domain"]),
            ("Tone",     post["tone"]),
            ("Hashtags", post["hashtags"]),
            ("CTA",      post["cta"]),
        ]
        for r, (label, value) in enumerate(detail_data):
            lc = detail_table.rows[r].cells[0]
            vc = detail_table.rows[r].cells[1]
            set_cell_bg(lc, "E8F8F7")
            lc.paragraphs[0].clear()
            add_run(lc.paragraphs[0], label, bold=True, size=9, color=TEAL)
            vc.paragraphs[0].clear()
            add_run(vc.paragraphs[0], value, size=9, color=DARK)

        doc.add_paragraph()

        # Prompt box
        prompt_label = doc.add_paragraph()
        add_run(prompt_label, "✍  Writing Prompt", bold=True, size=11, color=TEAL)

        prompt_box = doc.add_table(rows=1, cols=1)
        prompt_box.style = "Table Grid"
        pcell = prompt_box.rows[0].cells[0]
        set_cell_bg(pcell, "FFFDF0")
        pcell.paragraphs[0].clear()
        pcell.paragraphs[0].paragraph_format.space_before = Pt(4)
        pcell.paragraphs[0].paragraph_format.space_after  = Pt(4)
        add_run(pcell.paragraphs[0], post["prompt"], italic=True, size=10, color=DARK)

        if idx < len(WEEK1) - 1:
            doc.add_paragraph()
            divider(doc)
            doc.add_paragraph()
        else:
            doc.add_paragraph()

    # ── Footer ────────────────────────────────
    page_break(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_p, "LifeGate  ·  CEO LinkedIn Content Calendar  ·  Week 1 of 6",
            size=9, color=MGREY)
    add_run(footer_p, "\nGenerated: May 2026  ·  Confidential — Internal Use Only",
            size=9, color=MGREY)

    out_path = os.path.join(OUTPUT_DIR, "lifegate_linkedin_calendar_week1.docx")
    doc.save(out_path)
    print(f"✅  Calendar document saved → {out_path}")
    return out_path


# ──────────────────────────────────────────────
# DOCUMENT 2 – READY-TO-PUBLISH POSTS
# ──────────────────────────────────────────────

def build_posts_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_para.paragraph_format.space_before = Pt(40)
    add_run(cover_para, "LifeGate", bold=True, size=36, color=TEAL)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, "CEO LinkedIn Posts — Ready to Publish", bold=True, size=22, color=DARK)

    week_line = doc.add_paragraph()
    week_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(week_line, "WEEK 1  ·  June 1 – June 5, 2026", bold=False, size=14, color=MGREY)

    divider(doc)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tagline, '"Your trusted digital health companion."', italic=True, size=12, color=MGREY)

    doc.add_paragraph()

    instructions = doc.add_table(rows=1, cols=1)
    instructions.style = "Table Grid"
    ic = instructions.rows[0].cells[0]
    set_cell_bg(ic, "E8F8F7")
    ic.paragraphs[0].clear()
    add_run(ic.paragraphs[0],
            "📋  Instructions: Each post below is ready to copy and paste directly into LinkedIn. "
            "Review for any last-minute personalisation (e.g., adding a specific date, tagging a "
            "partner, or adding a photo/carousel). Posts are optimised for LinkedIn's algorithm: "
            "hook in line 1, value in the body, strong CTA at the close.",
            size=10, color=DARK)

    page_break(doc)

    # ── Individual posts ──────────────────────
    for idx, post in enumerate(WEEK1):
        # Header
        hdr_table = doc.add_table(rows=1, cols=1)
        hdr_table.style = "Table Grid"
        hc = hdr_table.rows[0].cells[0]
        set_cell_bg(hc, "0AADA2")
        hc.paragraphs[0].clear()
        add_run(hc.paragraphs[0],
                f"POST {post['post_number']}  ·  {post['day']}  ·  {post['theme']}",
                bold=True, color=WHITE, size=12)

        doc.add_paragraph()

        # Post body — preserve line breaks
        lines = post["post"].split("\n")
        for line in lines:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after  = Pt(2)
            add_run(lp, line, size=11, color=DARK)

        doc.add_paragraph()

        # Meta strip
        meta_strip = doc.add_table(rows=2, cols=2)
        meta_strip.style = "Table Grid"
        meta_vals = [
            ("Hashtags", post["hashtags"]),
            ("CTA",      post["cta"]),
        ]
        for r, (label, value) in enumerate(meta_vals):
            lc = meta_strip.rows[r].cells[0]
            vc = meta_strip.rows[r].cells[1]
            set_cell_bg(lc, "1A1A2E")
            lc.paragraphs[0].clear()
            add_run(lc.paragraphs[0], label, bold=True, size=9, color=WHITE)
            vc.paragraphs[0].clear()
            add_run(vc.paragraphs[0], value, size=9, color=DARK)

        if idx < len(WEEK1) - 1:
            doc.add_paragraph()
            divider(doc)
            doc.add_paragraph()
        else:
            doc.add_paragraph()

    # ── Footer ────────────────────────────────
    page_break(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_p, "LifeGate  ·  CEO LinkedIn Posts  ·  Week 1 of 6",
            size=9, color=MGREY)
    add_run(footer_p, "\nGenerated: May 2026  ·  Confidential — Internal Use Only",
            size=9, color=MGREY)

    out_path = os.path.join(OUTPUT_DIR, "lifegate_linkedin_posts_week1.docx")
    doc.save(out_path)
    print(f"✅  Posts document saved → {out_path}")
    return out_path


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────
if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    cal_path  = build_calendar_doc()
    post_path = build_posts_doc()
    print("\n📁  Both documents are available at:")
    print(f"   Calendar : /docs/{os.path.basename(cal_path)}")
    print(f"   Posts    : /docs/{os.path.basename(post_path)}")
