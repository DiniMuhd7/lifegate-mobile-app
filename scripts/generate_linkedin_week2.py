"""
LifeGate CEO LinkedIn Post Calendar – Week 2 Generator
Produces two Word documents:
  1. lifegate_linkedin_calendar_week2.docx  – calendar + prompts
  2. lifegate_linkedin_posts_week2.docx     – ready-to-publish posts

Week 2 domains:
  Mon Jun 8  – Health Metrics & Wearables
  Tue Jun 9  – LifeCoins & Accessible Payments
  Wed Jun 10 – Real-time Communication (IM + Voice/Video)
  Thu Jun 11 – Public Health Surveillance & Analytics
  Fri Jun 12 – Ambassador Campaign
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Brand colours ──────────────────────────────────────────────────────────
TEAL  = RGBColor(0x0A, 0xAD, 0xA2)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGREY = RGBColor(0xF5, 0xF7, 0xFA)
MGREY = RGBColor(0x64, 0x74, 0x8B)

OUTPUT_DIR = "/workspaces/lifegate-mobile-app/public/docs"

# ── Week 2 posts ───────────────────────────────────────────────────────────
WEEK2 = [
    {
        "day": "Monday, June 8, 2026",
        "post_number": "06",
        "theme": "Health Metrics & Wearables — Know Your Numbers",
        "domain": "Health Metrics / Wearable Sync",
        "tone": "Empowering · Data-driven",
        "hashtags": "#HealthMetrics #WearableTech #PreventiveHealth #LifeGate #DigitalHealth",
        "cta": "What health metric do you track daily? Tell us in the comments.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate on the power of continuous health "
            "monitoring. Explain that LifeGate syncs with wearables to track steps, active minutes, "
            "distance, and calories — and feeds this data into the patient's health profile so "
            "physicians have full context during case reviews. Use a compelling analogy: health data "
            "is like a financial statement — you can't manage what you don't measure. Highlight how "
            "passive data collection (vs manual entry) removes friction and encourages consistent "
            "health tracking. Tie it to early detection — small data trends that prevent big health "
            "surprises. Tone: empowering, data-driven, accessible. Length: 180–220 words."
        ),
        "post": (
            "You wouldn't manage a business without looking at the numbers.\n\n"
            "So why do most people go months — sometimes years — without looking at their "
            "health numbers?\n\n"
            "Steps taken. Active minutes. Calories burned. Distance covered. These aren't "
            "vanity metrics. They are early signals. Patterns that, when tracked consistently, "
            "can flag declining fitness, rising sedentary behaviour, or the early onset of "
            "metabolic conditions — before symptoms ever show up.\n\n"
            "LifeGate syncs seamlessly with your wearable device to capture these signals "
            "automatically. No manual entry. No friction. Just your data, working for you.\n\n"
            "But here's what makes our approach different:\n\n"
            "When you see a physician on LifeGate, they don't just read your symptoms report. "
            "They see your full health picture — your activity trends, your baseline metrics, "
            "your history. That context changes clinical decisions.\n\n"
            "A physician who knows you've been sedentary for three weeks, combined with your "
            "reported chest tightness, reaches a very different — and more accurate — conclusion "
            "than one who sees symptoms alone.\n\n"
            "That's not just smarter healthcare. That's safer healthcare.\n\n"
            "Know your numbers. Your body has been trying to tell you something.\n\n"
            "💬 What health metric do you track daily? Tell us in the comments.\n\n"
            "#HealthMetrics #WearableTech #PreventiveHealth #LifeGate #DigitalHealth"
        ),
    },
    {
        "day": "Tuesday, June 9, 2026",
        "post_number": "07",
        "theme": "LifeCoins — Making Quality Healthcare Affordable",
        "domain": "Payments / LifeCoins / Paystack",
        "tone": "Inclusive · Bold",
        "hashtags": "#HealthcareAffordability #LifeCoins #FinTech #LifeGate #AccessibleHealth",
        "cta": "Healthcare shouldn't be a luxury. Share this if you agree.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate introducing the LifeCoins credit "
            "system — the platform's approach to making quality healthcare affordable and flexible. "
            "Explain that patients can pay per consultation using LifeCoins or subscribe to a tier "
            "for unlimited access. Mention Paystack integration for seamless NGN/USD payments. "
            "Draw a powerful contrast: the cost of one hospital visit in Nigeria vs the cost of "
            "a LifeGate subscription that covers unlimited AI triage + physician validations. "
            "Frame affordability as a justice issue, not just a product feature. Tone: bold, "
            "inclusive, mission-driven. Length: 180–220 words."
        ),
        "post": (
            "In Nigeria, a single specialist consultation can cost between ₦10,000 and ₦50,000 "
            "— before tests, before transport, before lost income from taking the day off.\n\n"
            "For many families, that's not a healthcare decision. That's a financial crisis.\n\n"
            "We built LifeCoins because we believe cost should never be the reason someone "
            "avoids getting a health question answered.\n\n"
            "Here's how it works:\n\n"
            "🪙 Pay-per-use: Buy LifeCoins and spend them only when you need a consultation\n"
            "📦 Subscription tiers: Unlock unlimited AI triage + physician validations for a "
            "flat monthly fee\n"
            "💳 Powered by Paystack: Seamless, secure NGN and USD payments\n\n"
            "We designed LifeCoins to fit every budget — whether you're a student checking on "
            "a recurring headache or a family managing a chronic condition across multiple members.\n\n"
            "The question we ask ourselves every day at LifeGate is this:\n\n"
            "If someone in Maiduguri or Owerri or Yenagoa has a health concern at 11pm on a "
            "Sunday — can they get a credible, physician-backed answer without breaking the bank?\n\n"
            "The answer has to be yes.\n\n"
            "Healthcare shouldn't be a luxury. Share this if you agree. 👇\n\n"
            "#HealthcareAffordability #LifeCoins #FinTech #LifeGate #AccessibleHealth"
        ),
    },
    {
        "day": "Wednesday, June 10, 2026",
        "post_number": "08",
        "theme": "Real-time Communication — Closing the Gap Between Patient & Physician",
        "domain": "Instant Messaging / Voice & Video Calls (WebRTC)",
        "tone": "Human · Warm",
        "hashtags": "#Telemedicine #PatientPhysicianCommunication #LifeGate #HealthTech #WebRTC",
        "cta": "Tag a doctor or patient advocate who needs to see this.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate about the platform's real-time "
            "communication layer — instant messaging, voice calls, and video consultations "
            "between patients and physicians, all within the app. Paint a vivid before/after "
            "picture: before LifeGate, a patient waiting weeks for a follow-up appointment; "
            "after LifeGate, a physician pinging a patient directly through the app to clarify "
            "a result or check in after a prescription. Emphasise that this isn't telehealth "
            "as a gimmick — it's a clinical necessity for follow-up, relationship continuity, "
            "and reducing patient anxiety. Mention WebRTC-powered calls as secure and low-latency. "
            "Tone: human, warm, relational. Length: 190–230 words."
        ),
        "post": (
            "The hardest part of a diagnosis isn't always the diagnosis itself.\n\n"
            "It's the silence after.\n\n"
            "The patient who goes home with a report and a hundred unanswered questions. "
            "The one who calls the clinic and sits on hold. The one who turns to Google at "
            "2am because they didn't want to bother the doctor.\n\n"
            "We built LifeGate's communication layer to end that silence.\n\n"
            "Within the same app where a patient receives their physician-validated diagnosis, "
            "they can:\n\n"
            "💬 Message their physician directly — ask follow-up questions, share new symptoms, "
            "get clarification on a recommendation\n"
            "📞 Start a voice call — when typing isn't enough\n"
            "🎥 Join a video consultation — face-to-face, without leaving home\n\n"
            "All powered by WebRTC — secure, real-time, low-latency. No third-party apps. "
            "No switching between platforms. Everything in one place.\n\n"
            "For physicians, this changes the quality of care they can deliver. They're not "
            "just signing off on a case and moving on — they can follow up, check in, and "
            "build a real therapeutic relationship with their patients.\n\n"
            "Healthcare is not just about the moment of diagnosis. It's about everything that "
            "comes after.\n\n"
            "Tag a doctor or patient advocate who needs to see this. 👇\n\n"
            "#Telemedicine #PatientPhysicianCommunication #LifeGate #HealthTech #WebRTC"
        ),
    },
    {
        "day": "Thursday, June 11, 2026",
        "post_number": "09",
        "theme": "Public Health Surveillance — Data That Saves Communities",
        "domain": "Surveillance / Public Health Analytics / Research",
        "tone": "Authoritative · Purposeful",
        "hashtags": "#PublicHealth #HealthSurveillance #DataForGood #LifeGate #Nigeria",
        "cta": "Are you a researcher or public health official? Request API access at dshub.com.ng.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate on the platform's public health "
            "surveillance capability. Explain that LifeGate anonymises and aggregates clinical "
            "data from across Nigeria into a real-time epidemiological feed — covering all 36 "
            "states. This data is available through a public API for government agencies, NGOs, "
            "hospitals, and researchers. Highlight a specific use case: how early surveillance "
            "data can detect outbreak patterns weeks before traditional reporting systems catch "
            "them. Emphasise NDPA compliance — individual identity is never exposed. Position "
            "LifeGate as infrastructure for Nigeria's public health system, not just a consumer "
            "app. Tone: authoritative, purposeful, nationally minded. Length: 180–220 words."
        ),
        "post": (
            "What if Nigeria had a real-time map of where people are getting sick — before "
            "outbreaks become crises?\n\n"
            "That's not a future ambition. That's what LifeGate is building right now.\n\n"
            "Every interaction on our platform — every symptom reported, every AI triage "
            "completed, every physician-validated diagnosis — contributes to an anonymised "
            "epidemiological dataset that spans all 36 states.\n\n"
            "We aggregate that data into Nigeria's first physician-validated public health feed "
            "and make it available via a public API to:\n\n"
            "🏛️ Government health agencies\n"
            "🔬 Researchers and universities\n"
            "🏥 Hospitals and healthcare networks\n"
            "🌍 NGOs and international health partners\n\n"
            "The signal value is enormous. Traditional disease reporting systems lag by weeks. "
            "LifeGate's surveillance data moves in near real-time — giving public health "
            "officials the early warning they need to mobilise resources before a cluster "
            "becomes a crisis.\n\n"
            "And individual privacy is never compromised. Every data point is anonymised and "
            "NDPA 2023 compliant. No names. No identifiers. Just patterns that save communities.\n\n"
            "LifeGate isn't just a health app. It's public health infrastructure.\n\n"
            "Are you a researcher or public health official? Request API access at dshub.com.ng 👇\n\n"
            "#PublicHealth #HealthSurveillance #DataForGood #LifeGate #Nigeria"
        ),
    },
    {
        "day": "Friday, June 12, 2026",
        "post_number": "10",
        "theme": "Ambassador Campaign — Become a LifeGate Health Champion",
        "domain": "Ambassador Campaign / Community Advocacy",
        "tone": "Energetic · Community-driven",
        "hashtags": "#LifeGateAmbassador #HealthChampion #CommunityHealth #LifeGate #Nigeria",
        "cta": "DM 'AMBASSADOR' to learn how to join the LifeGate Health Champions network.",
        "prompt": (
            "Write a LinkedIn post from the CEO of LifeGate launching the LifeGate Health "
            "Champions Ambassador Programme. The programme invites healthcare professionals, "
            "community leaders, social media influencers, and patient advocates to represent "
            "LifeGate in their communities. Describe what ambassadors do: educate their "
            "communities about AI-powered health access, share LifeGate's mission, refer users "
            "through personalised referral codes, and host community health awareness sessions. "
            "Highlight what ambassadors receive: a personalised referral dashboard, recognition "
            "on the LifeGate platform, access to exclusive content, and the satisfaction of "
            "driving real health impact. Use energetic, rally-cry language. "
            "Tone: energetic, inclusive, community-powered. Length: 190–230 words."
        ),
        "post": (
            "Healthcare doesn't change from the top down. It changes community by community, "
            "conversation by conversation.\n\n"
            "That's why today, I'm officially launching the LifeGate Health Champions "
            "Ambassador Programme. 🚀\n\n"
            "We're looking for passionate individuals to carry the LifeGate mission into their "
            "communities — doctors, nurses, community health workers, students, influencers, "
            "patient advocates, and anyone who believes that better healthcare access is a "
            "right, not a privilege.\n\n"
            "As a LifeGate Health Champion, you will:\n\n"
            "📢 Educate your community about AI-assisted health access\n"
            "🔗 Share your personalised referral code and track your impact on a dedicated "
            "dashboard\n"
            "🎙️ Host community health awareness sessions (virtual or in-person)\n"
            "🌟 Get featured on the LifeGate platform and social channels\n"
            "🎁 Access exclusive content, early feature previews, and Champion-only resources\n\n"
            "You don't need to be a doctor to change healthcare in Nigeria. You just need to "
            "care enough to start a conversation.\n\n"
            "Every referral you make is a person who gets access to physician-validated health "
            "guidance — possibly for the first time in their life.\n\n"
            "That's not a small thing. That is legacy.\n\n"
            "DM me 'AMBASSADOR' to join the LifeGate Health Champions network. 💙\n\n"
            "#LifeGateAmbassador #HealthChampion #CommunityHealth #LifeGate #Nigeria"
        ),
    },
]


# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=11, color=None, font_name="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1, color=DARK, size=None, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sizes = {1: 22, 2: 16, 3: 13, 4: 11}
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p

def body(doc, text, italic=False, color=DARK, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p

def divider(doc, color_hex="0AADA2"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def page_break(doc):
    doc.add_page_break()

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


# ── Document 1: Calendar & Prompts ─────────────────────────────────────────

def build_calendar_doc():
    doc = Document()
    set_margins(doc)

    # Cover
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(40)
    add_run(cp, "LifeGate", bold=True, size=36, color=TEAL)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp, "CEO LinkedIn Post Calendar", bold=True, size=22, color=DARK)

    wl = doc.add_paragraph()
    wl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(wl, "WEEK 2  ·  June 8 – June 12, 2026", size=14, color=MGREY)

    divider(doc)

    tl = doc.add_paragraph()
    tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tl, '"Your trusted digital health companion."', italic=True, size=12, color=MGREY)

    doc.add_paragraph()

    mt = doc.add_table(rows=3, cols=2)
    mt.style = "Table Grid"
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate([
        ("Prepared for",   "CEO, LifeGate (DHSub)"),
        ("Campaign Period","Week 2 of 6 — June 2026"),
        ("Posts this week","5 posts (Mon – Fri)"),
    ]):
        lc = mt.rows[i].cells[0]; vc = mt.rows[i].cells[1]
        set_cell_bg(lc, "0AADA2")
        lc.paragraphs[0].clear(); add_run(lc.paragraphs[0], label, bold=True, color=WHITE, size=10)
        vc.paragraphs[0].clear(); add_run(vc.paragraphs[0], value, size=10, color=DARK)

    page_break(doc)

    # Intro
    heading(doc, "About This Calendar", level=2, color=TEAL)
    divider(doc)
    body(doc,
         "This document contains the Week 2 LinkedIn Post Calendar for the CEO of LifeGate. "
         "Week 2 expands into platform capabilities: Health Metrics & Wearables, LifeCoins "
         "affordable payments, Real-time Communication, Public Health Surveillance, and the "
         "LifeGate Health Champions Ambassador Campaign.")
    doc.add_paragraph()

    # Strategy table
    heading(doc, "Week 2 Content Strategy Overview", level=2, color=TEAL)
    divider(doc)

    st = doc.add_table(rows=6, cols=4)
    st.style = "Table Grid"
    hrow = st.rows[0]
    for j, h in enumerate(["Day", "Theme", "Domain", "Tone"]):
        cell = hrow.cells[j]
        set_cell_bg(cell, "0AADA2")
        cell.paragraphs[0].clear()
        add_run(cell.paragraphs[0], h, bold=True, color=WHITE, size=10)
    for i, post in enumerate(WEEK2):
        row = st.rows[i + 1]
        vals = [post["day"].split(",")[0], post["theme"], post["domain"], post["tone"]]
        for j, val in enumerate(vals):
            row.cells[j].paragraphs[0].clear()
            add_run(row.cells[j].paragraphs[0], val, size=9, color=DARK)
            if i % 2 == 0:
                set_cell_bg(row.cells[j], "F5F7FA")

    doc.add_paragraph()
    page_break(doc)

    # Post prompts
    heading(doc, "Week 2 — Post Prompts", level=1, color=TEAL, size=18, center=True)
    doc.add_paragraph()

    for idx, post in enumerate(WEEK2):
        ht = doc.add_table(rows=1, cols=1)
        ht.style = "Table Grid"
        hc = ht.rows[0].cells[0]
        set_cell_bg(hc, "0AADA2")
        hc.paragraphs[0].clear()
        add_run(hc.paragraphs[0], f"POST {post['post_number']}  ·  {post['day']}", bold=True, color=WHITE, size=12)

        dt = doc.add_table(rows=5, cols=2)
        dt.style = "Table Grid"
        for r, (label, value) in enumerate([
            ("Theme",    post["theme"]),
            ("Domain",   post["domain"]),
            ("Tone",     post["tone"]),
            ("Hashtags", post["hashtags"]),
            ("CTA",      post["cta"]),
        ]):
            lc = dt.rows[r].cells[0]; vc = dt.rows[r].cells[1]
            set_cell_bg(lc, "E8F8F7")
            lc.paragraphs[0].clear(); add_run(lc.paragraphs[0], label, bold=True, size=9, color=TEAL)
            vc.paragraphs[0].clear(); add_run(vc.paragraphs[0], value, size=9, color=DARK)

        doc.add_paragraph()
        pl = doc.add_paragraph()
        add_run(pl, "✍  Writing Prompt", bold=True, size=11, color=TEAL)

        pb = doc.add_table(rows=1, cols=1)
        pb.style = "Table Grid"
        pc = pb.rows[0].cells[0]
        set_cell_bg(pc, "FFFDF0")
        pc.paragraphs[0].clear()
        pc.paragraphs[0].paragraph_format.space_before = Pt(4)
        pc.paragraphs[0].paragraph_format.space_after  = Pt(4)
        add_run(pc.paragraphs[0], post["prompt"], italic=True, size=10, color=DARK)

        if idx < len(WEEK2) - 1:
            doc.add_paragraph(); divider(doc); doc.add_paragraph()
        else:
            doc.add_paragraph()

    page_break(doc)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, "LifeGate  ·  CEO LinkedIn Content Calendar  ·  Week 2 of 6\nGenerated: May 2026  ·  Confidential — Internal Use Only",
            size=9, color=MGREY)

    out = os.path.join(OUTPUT_DIR, "lifegate_linkedin_calendar_week2.docx")
    doc.save(out)
    print(f"✅  Calendar saved → {out}")
    return out


# ── Document 2: Ready-to-publish posts ─────────────────────────────────────

def build_posts_doc():
    doc = Document()
    set_margins(doc)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(40)
    add_run(cp, "LifeGate", bold=True, size=36, color=TEAL)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp, "CEO LinkedIn Posts — Ready to Publish", bold=True, size=22, color=DARK)

    wl = doc.add_paragraph()
    wl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(wl, "WEEK 2  ·  June 8 – June 12, 2026", size=14, color=MGREY)

    divider(doc)

    tl = doc.add_paragraph()
    tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(tl, '"Your trusted digital health companion."', italic=True, size=12, color=MGREY)

    doc.add_paragraph()

    it = doc.add_table(rows=1, cols=1)
    it.style = "Table Grid"
    ic = it.rows[0].cells[0]
    set_cell_bg(ic, "E8F8F7")
    ic.paragraphs[0].clear()
    add_run(ic.paragraphs[0],
            "📋  Instructions: Each post is ready to copy and paste into LinkedIn. "
            "Personalise where noted (tag a partner, add a photo, adjust a date). "
            "Posts are structured for LinkedIn's algorithm: strong hook line 1, "
            "value in the body, clear CTA at the close.",
            size=10, color=DARK)

    page_break(doc)

    for idx, post in enumerate(WEEK2):
        ht = doc.add_table(rows=1, cols=1)
        ht.style = "Table Grid"
        hc = ht.rows[0].cells[0]
        set_cell_bg(hc, "0AADA2")
        hc.paragraphs[0].clear()
        add_run(hc.paragraphs[0],
                f"POST {post['post_number']}  ·  {post['day']}  ·  {post['theme']}",
                bold=True, color=WHITE, size=12)

        doc.add_paragraph()

        for line in post["post"].split("\n"):
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after  = Pt(2)
            add_run(lp, line, size=11, color=DARK)

        doc.add_paragraph()

        ms = doc.add_table(rows=2, cols=2)
        ms.style = "Table Grid"
        for r, (label, value) in enumerate([("Hashtags", post["hashtags"]), ("CTA", post["cta"])]):
            lc = ms.rows[r].cells[0]; vc = ms.rows[r].cells[1]
            set_cell_bg(lc, "1A1A2E")
            lc.paragraphs[0].clear(); add_run(lc.paragraphs[0], label, bold=True, size=9, color=WHITE)
            vc.paragraphs[0].clear(); add_run(vc.paragraphs[0], value, size=9, color=DARK)

        if idx < len(WEEK2) - 1:
            doc.add_paragraph(); divider(doc); doc.add_paragraph()
        else:
            doc.add_paragraph()

    page_break(doc)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, "LifeGate  ·  CEO LinkedIn Posts  ·  Week 2 of 6\nGenerated: May 2026  ·  Confidential — Internal Use Only",
            size=9, color=MGREY)

    out = os.path.join(OUTPUT_DIR, "lifegate_linkedin_posts_week2.docx")
    doc.save(out)
    print(f"✅  Posts saved → {out}")
    return out


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    build_calendar_doc()
    build_posts_doc()
